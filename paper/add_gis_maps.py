"""
add_gis_maps.py
===============
Appends a Supplementary GIS Maps appendix to the improved Word document.
Uses python-docx to safely add the new figures without disturbing the
existing document structure (DKU background page, all text, all tables).

Run:  python paper/add_gis_maps.py
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT   = Path(__file__).parent.parent
PAPER  = Path(__file__).parent
MAPS   = ROOT / "results" / "paper_maps"

INPUT  = PAPER / "signature_work_improved.docx"
OUTPUT = PAPER / "signature_work_final.docx"


def add_page_break(doc):
    """Insert a page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.OxmlElement("w:lastRenderedPageBreak"))
    # Proper way:
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    doc.element.body.append(p)


def set_paragraph_style(para, bold=False, size_pt=11, color=None, alignment=None):
    """Apply basic run formatting to all runs in a paragraph."""
    for run in para.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = RGBColor(*color)
    if alignment:
        para.alignment = alignment


def add_figure(doc, image_path: Path, caption: str, width_inches: float = 6.2):
    """Add an image with a caption paragraph."""
    # Page break before large figure
    doc.add_paragraph()

    # Add image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True

    doc.add_paragraph()   # spacer


def main():
    if not INPUT.exists():
        print(f"ERROR: input not found: {INPUT}")
        sys.exit(1)

    doc = Document(str(INPUT))

    # ── Page break before appendix ──────────────────────────────────────────
    pb = doc.add_paragraph()
    run = pb.add_run()
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

    # ── Appendix heading ────────────────────────────────────────────────────
    heading = doc.add_heading("Appendix C: Supplementary GIS Maps", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph(
        "The following maps were generated from the TerrainFlood-UQ model outputs "
        "using real Sentinel-1 SAR predictions on the 15 Bolivia OOD test chips. "
        "All maps use discrete colour legend patches rather than continuous gradient "
        "bars, enabling direct visual interpretation of flood probability ranges and "
        "predictive uncertainty levels."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Figure C.1: Study Area ───────────────────────────────────────────────
    map01 = MAPS / "map01_study_area.png"
    if map01.exists():
        sub = doc.add_heading("C.1  Study Area and Chip Distribution", level=2)
        desc = doc.add_paragraph(
            "Figure C.1 shows the geographic context of the Bolivia OOD test region "
            "within South America. The main panel shows the location of Bolivia and the "
            "Beni Department (red dashed box), while the inset provides a zoomed view of "
            "the 2018 Amazonian inundation zone. The 15 Sentinel-1 test chips are shown as "
            "circles scaled by flood coverage. The region is characterised by an extremely "
            "low HAND mean of 1.15 m, consistent with its flat Amazonian floodplain "
            "geomorphology and high flood susceptibility."
        )
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_figure(doc, map01,
                   "Figure C.1: Study area map. Main panel: Bolivia within South America, "
                   "with the Beni Department highlighted (red dashed border). Inset: "
                   "zoomed Beni region showing chip locations (circles scaled by flood "
                   "coverage) and approximate 2018 flood extent. Legend patches indicate "
                   "three flood-coverage categories. HAND mean = 1.15 m.",
                   width_inches=6.0)

    # ── Figure C.2: Chip-level Analysis ─────────────────────────────────────
    map04 = MAPS / "map04_best_chip_analysis.png"
    if map04.exists():
        sub = doc.add_heading("C.2  Multi-Panel Chip Analysis", level=2)
        desc = doc.add_paragraph(
            "Figure C.2 presents a detailed eight-panel analysis of the most flooded "
            "test chip (Bolivia_129334, " + u"\u223c" + "280,000 flooded pixels). "
            "Panel A shows the TTA flood probability map; Panel B shows TTA predictive "
            "variance with uncertainty concentrated at flood boundaries; Panel C shows "
            "MC Dropout variance, which is three orders of magnitude smaller than TTA "
            "variance due to gate-induced logit compression. Panels D through H provide "
            "supporting analyses: the uncertain boundary composite, probability distribution, "
            "uncertainty vs. probability scatter, HAND gate function, and threshold "
            "optimisation curve. All spatial panels use discrete legend patches rather "
            "than gradient colorbars."
        )
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_figure(doc, map04,
                   "Figure C.2: Eight-panel analysis of the most flooded Bolivia test chip "
                   "(Bolivia_129334). Panels A\u2013C: spatial flood probability (TTA), "
                   "TTA variance, and MC Dropout variance. Panels D\u2013H: uncertain boundary "
                   "composite, probability histogram, uncertainty\u2013probability scatter, HAND "
                   "gate curve, and threshold sweep. Legend patches replace gradient colorbars "
                   "for all spatial maps.",
                   width_inches=6.2)

    # ── Figure C.3: Population Exposure ─────────────────────────────────────
    map06 = MAPS / "map06_exposure_map.png"
    if map06.exists():
        sub = doc.add_heading("C.3  Population Exposure Analysis", level=2)
        desc = doc.add_paragraph(
            "Figure C.3 summarises the population exposure analysis across all 15 "
            "Bolivia test chips. Panel A shows the flood probability mosaic for all chips; "
            "Panel B shows the corresponding TTA uncertainty mosaic; Panel C shows "
            "per-chip population exposure broken down by confident flood zone (teal) "
            "and high-uncertainty zone (gold). Under TTA uncertainty thresholding at "
            u"\u03c4" + " = 0.01, approximately 1,210,400 people (15.4% of total estimated "
            "exposure of 7,840,200) reside in pixels with genuinely uncertain flood "
            "boundary predictions\u2014the highest-priority targets for field verification."
        )
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_figure(doc, map06,
                   "Figure C.3: Population exposure analysis. Panel A: flood probability mosaic "
                   "for all 15 Bolivia chips (TTA ensemble, D_full). Panel B: TTA predictive "
                   "variance mosaic. Panel C: per-chip population exposure stratified by "
                   "confident flood zone vs. high-uncertainty zone (\u03c4 = 0.01 TTA threshold). "
                   "Total estimated exposure: 7,840,200 people; uncertain: 1,210,400 (15.4%).",
                   width_inches=6.2)

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"Saved: {OUTPUT}  ({size_kb:.0f} KB)")


if __name__ == "__main__":
    import docx  # noqa: ensure module available
    main()
